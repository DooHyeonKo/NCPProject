from io import BytesIO
from pathlib import Path

import fitz
from docx import Document as DocxDocument


def extract_text_from_txt(file_path: str) -> str:
    path = Path(file_path)
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="ignore")


def extract_text_from_pdf(file_path: str) -> str:
    contents = []
    with fitz.open(file_path) as pdf:
        for page in pdf:
            contents.append(page.get_text("text"))
    return "\n".join(contents).strip()


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def extract_text(file_path: str, file_type: str) -> str:
    file_type = file_type.lower()
    if file_type == "txt":
        return extract_text_from_txt(file_path)
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    if file_type == "docx":
        return extract_text_from_docx(file_path)
    return ""


def extract_text_from_txt_bytes(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="ignore")


def extract_text_from_pdf_bytes(content: bytes) -> str:
    contents = []
    with fitz.open(stream=content, filetype="pdf") as pdf:
        for page in pdf:
            contents.append(page.get_text("text"))
    return "\n".join(contents).strip()


def extract_text_from_docx_bytes(content: bytes) -> str:
    doc = DocxDocument(BytesIO(content))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def extract_text_from_bytes(content: bytes, file_type: str) -> str:
    file_type = file_type.lower()
    if file_type == "txt":
        return extract_text_from_txt_bytes(content)
    if file_type == "pdf":
        return extract_text_from_pdf_bytes(content)
    if file_type == "docx":
        return extract_text_from_docx_bytes(content)
    return ""
